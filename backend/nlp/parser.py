"""
SkillPathForge AI — File Parser
---------------------------------
Extracts raw text from uploaded resume files.

Supports:
  - PDF files  (.pdf)  via PyMuPDF
  - Word files (.docx) via python-docx

Returns:
  - Raw text string
  - Detected sections (Experience, Skills, Education etc.)
  - Basic metadata (file type, page count)

No NLP happens here — just clean text extraction.
Text is passed to extractor.py for skill extraction.
"""

import io
import re
from pathlib import Path
from typing import Optional

import fitz                          # PyMuPDF
from docx import Document            # python-docx


# ── Section header patterns ───────────────────────────────────────
# Common resume section headers to detect
# Used to split resume into sections for confidence scoring

SECTION_PATTERNS: list[str] = [
    "work experience",
    "professional experience",
    "employment history",
    "experience",
    "skills",
    "technical skills",
    "core competencies",
    "education",
    "academic background",
    "certifications",
    "certificates",
    "achievements",
    "accomplishments",
    "projects",
    "personal projects",
    "summary",
    "professional summary",
    "objective",
    "training",
    "courses",
    "coursework",
    "volunteer",
    "publications",
    "awards",
    "languages",
    "interests",
    "hobbies",
]


# ── Supported file types ──────────────────────────────────────────

SUPPORTED_EXTENSIONS: set[str] = {".pdf", ".docx", ".doc"}


# ── ParsedResume dataclass ────────────────────────────────────────

class ParsedResume:
    """
    Container for parsed resume data.
    Passed to extractor.py for NLP processing.
    """

    def __init__(
        self,
        raw_text:   str,
        sections:   dict[str, str],
        metadata:   dict,
        file_type:  str,
    ):
        self.raw_text  = raw_text    # full cleaned text
        self.sections  = sections    # {section_name: section_text}
        self.metadata  = metadata    # file info
        self.file_type = file_type   # "pdf" or "docx"

    def get_section(self, name: str) -> str:
        """
        Returns text for a specific section.
        Case-insensitive lookup.

        Args:
            name: section name to look up

        Returns:
            section text or empty string
        """
        name_lower = name.lower()
        for key, value in self.sections.items():
            if name_lower in key.lower():
                return value
        return ""

    def has_section(self, name: str) -> bool:
        """
        Checks if a section exists in the resume.

        Args:
            name: section name to check

        Returns:
            True if section found
        """
        return bool(self.get_section(name))

    def to_dict(self) -> dict:
        """
        Converts to dict for API response.
        """
        return {
            "raw_text":  self.raw_text,
            "sections":  self.sections,
            "metadata":  self.metadata,
            "file_type": self.file_type,
        }


# ── PDF Parser ────────────────────────────────────────────────────

def parse_pdf(file_bytes: bytes) -> ParsedResume:
    """
    Extracts text from PDF file using PyMuPDF.

    Handles:
      - Multi-page PDFs
      - Encoding artifacts (Â, â€™ etc.)
      - Extra whitespace and blank lines
      - Section detection

    Args:
        file_bytes: raw PDF file bytes

    Returns:
        ParsedResume object with text and sections
    """
    try:
        # Open PDF from bytes
        pdf_document = fitz.open(
            stream=file_bytes,
            filetype="pdf"
        )

        pages_text: list[str] = []
        page_count = len(pdf_document)

        # Extract text from each page
        for page_num in range(page_count):
            page = pdf_document[page_num]
            page_text = page.get_text("text")
            if page_text.strip():
                pages_text.append(page_text)

        pdf_document.close()

        # Combine all pages
        raw_text = "\n".join(pages_text)

        # Clean the text
        cleaned_text = _clean_text(raw_text)

        # Detect sections
        sections = _detect_sections(cleaned_text)

        metadata = {
            "file_type":  "pdf",
            "page_count": page_count,
            "char_count": len(cleaned_text),
            "word_count": len(cleaned_text.split()),
        }

        return ParsedResume(
            raw_text=cleaned_text,
            sections=sections,
            metadata=metadata,
            file_type="pdf",
        )

    except Exception as e:
        raise ValueError(f"PDF parsing failed: {str(e)}")


# ── DOCX Parser ───────────────────────────────────────────────────

def parse_docx(file_bytes: bytes) -> ParsedResume:
    """
    Extracts text from Word document using python-docx.

    Handles:
      - Paragraphs and headings
      - Tables (common in resumes)
      - Section detection

    Args:
        file_bytes: raw DOCX file bytes

    Returns:
        ParsedResume object with text and sections
    """
    try:
        # Open DOCX from bytes
        doc = Document(io.BytesIO(file_bytes))

        paragraphs: list[str] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in paragraphs:
                        paragraphs.append(text)

        # Combine all text
        raw_text = "\n".join(paragraphs)

        # Clean the text
        cleaned_text = _clean_text(raw_text)

        # Detect sections
        sections = _detect_sections(cleaned_text)

        metadata = {
            "file_type":       "docx",
            "paragraph_count": len(doc.paragraphs),
            "table_count":     len(doc.tables),
            "char_count":      len(cleaned_text),
            "word_count":      len(cleaned_text.split()),
        }

        return ParsedResume(
            raw_text=cleaned_text,
            sections=sections,
            metadata=metadata,
            file_type="docx",
        )

    except Exception as e:
        raise ValueError(f"DOCX parsing failed: {str(e)}")


# ── Main parse function ───────────────────────────────────────────

def parse_file(
    file_bytes: bytes,
    filename:   str,
) -> ParsedResume:
    """
    Main entry point — routes to correct parser
    based on file extension.

    Args:
        file_bytes: raw file bytes from upload
        filename:   original filename with extension

    Returns:
        ParsedResume object

    Raises:
        ValueError: if file type not supported
    """
    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: {extension}. "
            f"Supported: {SUPPORTED_EXTENSIONS}"
        )

    if extension == ".pdf":
        return parse_pdf(file_bytes)

    elif extension in {".docx", ".doc"}:
        return parse_docx(file_bytes)

    else:
        raise ValueError(f"Cannot parse file: {filename}")


# ── Text cleaning ─────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """
    Cleans raw extracted text.

    Removes:
      - Encoding artifacts (Â, â€™, Ã etc.)
      - Excessive whitespace
      - Repeated newlines
      - Non-printable characters

    Preserves:
      - Line structure (important for section detection)
      - Punctuation
      - Numbers

    Args:
        text: raw extracted text

    Returns:
        cleaned text string
    """
    # Remove common encoding artifacts
    artifacts = [
        "Â", "â€™", "â€œ", "â€", "Ã©", "Ã ",
        "â€¦", "â€", "â€˜", "\x00", "\ufeff",
    ]
    for artifact in artifacts:
        text = text.replace(artifact, "")

    # Remove non-printable characters except newlines and tabs
    text = re.sub(r"[^\x20-\x7E\n\t]", " ", text)

    # Replace tabs with spaces
    text = text.replace("\t", " ")

    # Remove lines with only special characters
    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        # Keep line if it has at least 2 alphanumeric characters
        if len(re.findall(r"[a-zA-Z0-9]", stripped)) >= 2:
            cleaned_lines.append(stripped)

    # Rejoin with newlines
    text = "\n".join(cleaned_lines)

    # Remove excessive blank lines (max 2 consecutive)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Remove excessive spaces
    text = re.sub(r" {2,}", " ", text)

    return text.strip()


# ── Section detection ─────────────────────────────────────────────

def _detect_sections(text: str) -> dict[str, str]:
    """
    Splits resume text into named sections.

    Detects common section headers and splits
    text accordingly. Section text is then used
    by confidence scorer to determine context.

    Args:
        text: cleaned resume text

    Returns:
        dict mapping section_name → section_text

    Example:
        {
            "experience": "Built ML pipelines...",
            "skills": "Python, SQL, Docker...",
            "education": "BSc Computer Science..."
        }
    """
    sections: dict[str, str] = {}
    lines = text.split("\n")

    current_section = "header"
    current_lines: list[str] = []

    for line in lines:
        line_lower = line.lower().strip()

        # Check if this line is a section header
        matched_section = None
        for pattern in SECTION_PATTERNS:
            if (
                line_lower == pattern or
                line_lower.startswith(pattern) and
                len(line_lower) < len(pattern) + 10
            ):
                matched_section = pattern
                break

        if matched_section:
            # Save previous section
            if current_lines:
                sections[current_section] = "\n".join(
                    current_lines
                ).strip()

            # Start new section
            current_section = matched_section
            current_lines = []
        else:
            current_lines.append(line)

    # Save final section
    if current_lines:
        sections[current_section] = "\n".join(
            current_lines
        ).strip()

    # Always ensure these sections exist
    if "skills" not in sections:
        sections["skills"] = ""
    if "experience" not in sections:
        sections["experience"] = ""

    return sections


# ── Validate file ─────────────────────────────────────────────────

def validate_file(
    file_bytes: bytes,
    filename:   str,
    max_size_mb: float = 5.0,
) -> dict:
    """
    Validates uploaded file before parsing.

    Checks:
      - File size within limit
      - File extension supported
      - File not empty

    Args:
        file_bytes:  raw file bytes
        filename:    original filename
        max_size_mb: maximum file size in MB

    Returns:
        dict with valid: bool and message: str
    """
    # Check empty
    if not file_bytes:
        return {
            "valid":   False,
            "message": "File is empty"
        }

    # Check size
    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > max_size_mb:
        return {
            "valid":   False,
            "message": (
                f"File too large: {size_mb:.1f}MB. "
                f"Maximum: {max_size_mb}MB"
            )
        }

    # Check extension
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        return {
            "valid":   False,
            "message": (
                f"Unsupported file type: {extension}. "
                f"Please upload PDF or DOCX"
            )
        }

    return {
        "valid":   True,
        "message": "File is valid"
    }